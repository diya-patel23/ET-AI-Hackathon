import logging

logger = logging.getLogger("parsers")


def parse_pdf(path: str) -> str:
    import pdfplumber

    text_parts = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
                # tables often carry the real data in inspection/maintenance PDFs
                for table in page.extract_tables():
                    for row in table:
                        text_parts.append(" | ".join(c or "" for c in row))
    except Exception as e:
        logger.error("Failed to parse PDF %s: %s", path, e)
    return "\n".join(text_parts).strip()


def parse_docx(path: str) -> str:
    import docx

    try:
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts).strip()
    except Exception as e:
        logger.error("Failed to parse DOCX %s: %s", path, e)
        return ""


def parse_xlsx(path: str) -> str:
    import pandas as pd

    try:
        sheets = pd.read_excel(path, sheet_name=None)
        parts = []
        for name, df in sheets.items():
            parts.append(f"--- Sheet: {name} ---")
            parts.append(df.to_csv(index=False))
        return "\n".join(parts).strip()
    except Exception as e:
        logger.error("Failed to parse XLSX %s: %s", path, e)
        return ""


def parse_csv(path: str) -> str:
    import pandas as pd

    try:
        df = pd.read_csv(path)
        return df.to_csv(index=False).strip()
    except Exception as e:
        logger.error("Failed to parse CSV %s: %s", path, e)
        return ""


def parse_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as e:
        logger.error("Failed to parse TXT %s: %s", path, e)
        return ""


PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "xlsx": parse_xlsx,
    "csv": parse_csv,
    "txt": parse_txt,
    "email": parse_txt,  # .eml treated as plain text for the demo
}


def parse_document(path: str, file_type: str) -> str:
    parser = PARSERS.get(file_type)
    if not parser:
        return ""
    return parser(path)
